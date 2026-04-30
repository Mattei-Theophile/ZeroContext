import os
import io
from nlp import AnonymizeResponse
from mode.StrategyAnonymization import StrategyAnonymization
import pymupdf  # PyMuPDF
from dataclasses import dataclass
from docx import Document
from odf.opendocument import load
from odf import teletype, text

@dataclass
class FileAnonymizeResult:
    response: AnonymizeResponse
    file_bytes: bytes
    output_filename: str

class FileAnonymization(StrategyAnonymization):
    def anonymize(self, file_path: str) -> FileAnonymizeResult:
        print(f"Anonymizing file: {file_path}")
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".pdf":
            original_text = self.extract_text_from_pdf(file_path)
        elif ext == ".docx":
            original_text = self.extract_text_from_docx(file_path)
        elif ext == ".odt":
            original_text = self.extract_text_from_odt(file_path)
        elif ext == ".txt":
            original_text = self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        response = self.perform_anonymization(original_text)
        output_filename = f"anonymized_{os.path.basename(file_path)}"
        replacements = self.build_replacements(response.matchTable)

        if ext == ".pdf":
            file_bytes = self.redact_pdf(file_path, replacements)
        elif ext == ".docx":
            file_bytes = self.redact_docx(file_path, replacements)
        elif ext == ".odt":
            file_bytes = self.redact_odt(file_path, replacements)
        elif ext == ".txt":
            file_bytes = self.redact_txt(original_text, replacements)
        
        return FileAnonymizeResult(
            response=response,
            file_bytes=file_bytes,
            output_filename=output_filename
        )

    def placeholder_to_label(self, placeholder: str) -> str:
        if placeholder.startswith("[") and placeholder.endswith("]") and "_" in placeholder:
            inner = placeholder[1:-1]
            label = inner.rsplit("_", 1)[0]
            return f"[{label}]"
        return placeholder

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        doc = pymupdf.open(pdf_path)
        try:
            return "\n".join(page.get_text("text") for page in doc)
        finally:
            doc.close()

    def extract_text_from_docx(self, docx_path: str) -> str:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])

    def extract_text_from_odt(self, odt_path: str) -> str:
        textdoc = load(odt_path)
        all_paragraphs = textdoc.getElementsByType(text.P)
        return "\n".join([teletype.get_string(p) for p in all_paragraphs])

    def extract_text_from_txt(self, txt_path: str) -> str:
        with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def build_replacements(self, match_table: dict[str, str]) -> list[tuple[str, str]]:
        replacements: list[tuple[str, str]] = []
        for placeholder, original_text in match_table.items():
            replacement_label = self.placeholder_to_label(placeholder)
            if not original_text.strip():
                continue
            replacements.append((original_text, replacement_label))
        replacements.sort(key=lambda item: len(item[0]), reverse=True)
        return replacements

    def redact_pdf(self, input_pdf: str, replacements: list[tuple[str, str]]) -> bytes:
        doc = pymupdf.open(input_pdf)
        try:
            for page in doc:
                used_rects: list[pymupdf.Rect] = []
                for original_text, replacement_label in replacements:
                    rects = page.search_for(original_text)
                    for rect in rects:
                        if any(rect.intersects(existing) for existing in used_rects):
                            continue
                        page.add_redact_annot(
                            rect,
                            text=replacement_label,
                            fill=(0, 0, 0),
                            text_color=(1, 1, 1),
                        )
                        used_rects.append(rect)
                page.apply_redactions()
            return doc.tobytes()
        finally:
            doc.close()

    def redact_docx(self, input_docx: str, replacements: list[tuple[str, str]]) -> bytes:
        doc = Document(input_docx)
        for para in doc.paragraphs:
            for original, replacement in replacements:
                if original in para.text:
                    para.text = para.text.replace(original, replacement)
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for original, replacement in replacements:
                            if original in para.text:
                                para.text = para.text.replace(original, replacement)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def redact_odt(self, input_odt: str, replacements: list[tuple[str, str]]) -> bytes:
        textdoc = load(input_odt)
        all_paragraphs = textdoc.getElementsByType(text.P)
        for p in all_paragraphs:
            # This is a bit tricky with odfpy as it doesn't support simple text replacement well
            # if there's complex formatting. But for basic anonymization:
            current_text = teletype.get_string(p)
            new_text = current_text
            for original, replacement in replacements:
                new_text = new_text.replace(original, replacement)
            
            if new_text != current_text:
                # Clear children and add new text
                while p.firstChild:
                    p.removeChild(p.firstChild)
                p.addText(new_text)

        buffer = io.BytesIO()
        textdoc.save(buffer)
        return buffer.getvalue()

    def redact_txt(self, original_text: str, replacements: list[tuple[str, str]]) -> bytes:
        anonymized_text = original_text
        for original, replacement in replacements:
            anonymized_text = anonymized_text.replace(original, replacement)
        return anonymized_text.encode("utf-8")
